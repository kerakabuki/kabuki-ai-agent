# -*- coding: utf-8 -*-
"""docxのページ区切り(lastRenderedPageBreak)から各段落のページ番号を推定し、
表・テキストボックス内のテキストも確認する"""
import sys

import docx
from docx.oxml.ns import qn

PATH = r"D:\気良歌舞伎\台本\曽根崎心中\曽根崎心中.docx"


def main():
    d = docx.Document(PATH)
    page = 1
    page_of = {}
    for i, p in enumerate(d.paragraphs):
        if p._element.findall(".//" + qn("w:lastRenderedPageBreak")):
            page += 1
        for br in p._element.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                page += 1
        page_of[i] = page

    # ページ11-13の段落を表示
    for i, p in enumerate(d.paragraphs):
        if 11 <= page_of[i] <= 13 and p.text.strip():
            print(f"p{page_of[i]} #{i} | {p.text.strip()[:60]}")

    print("---- tables:", len(d.tables))
    for ti, tbl in enumerate(d.tables):
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    print(f"table{ti} | {t[:60]}")

    # テキストボックス
    body = d.element.body
    boxes = body.findall(".//" + qn("w:txbxContent"))
    print("---- textboxes:", len(boxes))
    for b in boxes:
        text = "".join(n.text or "" for n in b.iter(qn("w:t")))
        if text.strip():
            print("box |", text.strip()[:80])


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
